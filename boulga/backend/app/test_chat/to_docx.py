"""to_docx.py — Convertit un document JSON (grammaire compositionnelle) en .docx.

Robustesse : un bloc malformé, partiel ou de type inconnu est ignoré (warning),
jamais fatal. Un JSON tronqué produit un document avec les blocs valides.
"""

from __future__ import annotations

import io
import logging
from typing import Any, Optional

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

logger = logging.getLogger("boulga.to_docx")

_ALIGN = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# Tailles par défaut des titres par niveau (points)
_HEADING_SIZE = {1: 22, 2: 17, 3: 14, 4: 12}


def _hex_to_rgb(value: Optional[str]) -> Optional[RGBColor]:
    """Convertit '#RRGGBB' (ou 'RRGGBB') en RGBColor. None si invalide."""
    if not value or not isinstance(value, str):
        return None
    v = value.strip().lstrip("#")
    if len(v) != 6:
        return None
    try:
        return RGBColor(int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16))
    except ValueError:
        return None


def _shade(paragraph, hex_color: Optional[str]) -> None:
    """Applique un fond coloré (shading) à un paragraphe."""
    rgb = (hex_color or "").strip().lstrip("#")
    if len(rgb) != 6:
        return
    try:
        int(rgb, 16)
    except ValueError:
        return
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb.upper())
    pPr.append(shd)


def _cell_shade(cell, hex_color: Optional[str]) -> None:
    rgb = (hex_color or "").strip().lstrip("#")
    if len(rgb) != 6:
        return
    try:
        int(rgb, 16)
    except ValueError:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb.upper())
    tc_pr.append(shd)


def _apply_run_style(run, *, color=None, size=None, bold=None, italic=None) -> None:
    rgb = _hex_to_rgb(color)
    if rgb is not None:
        run.font.color.rgb = rgb
    if size:
        try:
            run.font.size = Pt(int(size))
        except (ValueError, TypeError):
            pass
    if bold is not None:
        run.font.bold = bool(bold)
    if italic is not None:
        run.font.italic = bool(italic)


def _item_to_parts(item: Any) -> dict:
    """Normalise un item de liste (str ou dict) en dict de propriétés."""
    if isinstance(item, str):
        return {"text": item}
    if isinstance(item, dict):
        return {
            "text": str(item.get("text", "")),
            "color": item.get("color"),
            "bold": item.get("bold"),
            "italic": item.get("italic"),
        }
    return {"text": str(item)}


# ── Rendu des blocs ──────────────────────────────────────────────────────────

def _render_heading(doc, b: dict) -> None:
    level = b.get("level", 1)
    try:
        level = max(1, min(int(level), 4))
    except (ValueError, TypeError):
        level = 1
    p = doc.add_paragraph()
    if b.get("align") in _ALIGN:
        p.alignment = _ALIGN[b["align"]]
    run = p.add_run(str(b.get("text", "")))
    run.font.bold = True
    size = b.get("size") or _HEADING_SIZE.get(level, 14)
    _apply_run_style(run, color=b.get("color"), size=size, bold=True)


def _render_paragraph(doc, b: dict) -> None:
    p = doc.add_paragraph()
    if b.get("align") in _ALIGN:
        p.alignment = _ALIGN[b["align"]]
    run = p.add_run(str(b.get("text", "")))
    _apply_run_style(
        run, color=b.get("color"), size=b.get("size"),
        bold=b.get("bold"), italic=b.get("italic"),
    )


def _render_bullet_list(doc, b: dict) -> None:
    bullet = b.get("bullet") or "•"
    if not isinstance(bullet, str) or not bullet:
        bullet = "•"
    color = b.get("color")
    size = b.get("size")
    for item in b.get("items", []):
        parts = _item_to_parts(item)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        run = p.add_run(f"{bullet}  {parts['text']}")
        _apply_run_style(
            run,
            color=parts.get("color") or color,
            size=size,
            bold=parts.get("bold"),
            italic=parts.get("italic"),
        )


def _render_numbered_list(doc, b: dict) -> None:
    color = b.get("color")
    size = b.get("size")
    for i, item in enumerate(b.get("items", []), start=1):
        parts = _item_to_parts(item)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        run = p.add_run(f"{i}.  {parts['text']}")
        _apply_run_style(
            run,
            color=parts.get("color") or color,
            size=size,
            bold=parts.get("bold"),
            italic=parts.get("italic"),
        )


def _render_table(doc, b: dict) -> None:
    headers = b.get("headers", []) or []
    rows = b.get("rows", []) or []
    col_count = max(len(headers), max((len(r) for r in rows), default=0))
    if col_count == 0:
        return

    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"

    header_bg = b.get("header_bg")
    header_color = b.get("header_color") or ("#FFFFFF" if header_bg else None)
    zebra = b.get("zebra")

    if headers:
        cells = table.add_row().cells
        for i in range(col_count):
            text = str(headers[i]) if i < len(headers) else ""
            cell = cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            run.font.bold = True
            rgb = _hex_to_rgb(header_color)
            if rgb is not None:
                run.font.color.rgb = rgb
            if header_bg:
                _cell_shade(cell, header_bg)

    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i in range(col_count):
            val = row[i] if i < len(row) else ""
            cells[i].text = str(val) if val is not None else ""
        if zebra and r_idx % 2 == 1:
            for c in cells:
                _cell_shade(c, zebra)


def _render_block(doc, b: dict) -> None:
    """Encadré à fond coloré. Rendu comme un paragraphe ombré (texte simple)
    ou, si 'content' est fourni, comme une suite de blocs sur fond coloré.
    """
    bg = b.get("bg")
    text_color = b.get("text_color")
    align = b.get("align")

    # Texte simple
    if b.get("text") is not None:
        p = doc.add_paragraph()
        if align in _ALIGN:
            p.alignment = _ALIGN[align]
        if bg:
            _shade(p, bg)
        run = p.add_run(str(b.get("text", "")))
        _apply_run_style(run, color=text_color)
        return

    # Contenu imbriqué : on applique le fond à chaque paragraphe enfant rendu
    content = b.get("content")
    if isinstance(content, list):
        for child in content:
            if not isinstance(child, dict):
                continue
            # On force la couleur de texte du bloc si l'enfant n'en a pas
            if text_color and "color" not in child and child.get("type") in (
                "paragraph", "heading", "bullet_list", "numbered_list",
            ):
                child = {**child, "color": text_color}
            before = len(doc.paragraphs)
            _render_block_dispatch(doc, child)
            # Appliquer le fond aux paragraphes nouvellement créés
            if bg:
                for p in doc.paragraphs[before:]:
                    _shade(p, bg)


def _render_divider(doc, b: dict) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str((b.get("thickness") or 1) * 6))
    bottom.set(qn("w:space"), "1")
    color = (b.get("color") or "999999").lstrip("#")
    bottom.set(qn("w:color"), color.upper() if len(color) == 6 else "999999")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _render_spacer(doc, b: dict) -> None:
    p = doc.add_paragraph()
    size = b.get("size") or 12
    try:
        p.paragraph_format.space_after = Pt(int(size))
    except (ValueError, TypeError):
        p.paragraph_format.space_after = Pt(12)


def _render_page_break(doc, b: dict) -> None:
    doc.add_page_break()


_DISPATCH = {
    "heading": _render_heading,
    "paragraph": _render_paragraph,
    "bullet_list": _render_bullet_list,
    "numbered_list": _render_numbered_list,
    "table": _render_table,
    "block": _render_block,
    "divider": _render_divider,
    "spacer": _render_spacer,
    "page_break": _render_page_break,
}


def _render_block_dispatch(doc, b: dict) -> None:
    btype = b.get("type")
    handler = _DISPATCH.get(btype)
    if handler is None:
        logger.warning("Bloc de type inconnu ignoré: %r", btype)
        return
    try:
        handler(doc, b)
    except Exception as exc:  # un bloc ne doit jamais casser le document
        logger.warning("Bloc %r ignoré (erreur de rendu): %s", btype, exc)


def render_docx(document: dict) -> bytes:
    """Construit un .docx depuis le JSON et retourne ses bytes."""
    doc = DocxDocument()

    # Marges raisonnables
    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(64)
        section.right_margin = Pt(64)

    title = document.get("title")
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(title))
        run.font.bold = True
        run.font.size = Pt(26)

    blocks = document.get("blocks", [])
    if not isinstance(blocks, list):
        blocks = []

    for b in blocks:
        if isinstance(b, dict):
            _render_block_dispatch(doc, b)
        else:
            logger.warning("Bloc non-dict ignoré: %r", type(b))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()