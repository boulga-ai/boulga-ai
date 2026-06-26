from __future__ import annotations

import logging
from io import BytesIO
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

logger = logging.getLogger(__name__)


def _safe_color(hex_color: Any) -> tuple[int, int, int] | None:
    if not isinstance(hex_color, str):
        return None
    color = hex_color.strip().lstrip("#")
    if len(color) not in {3, 6}:
        return None
    if len(color) == 3:
        color = "".join([c * 2 for c in color])
    try:
        return tuple(int(color[i : i + 2], 16) for i in (0, 2, 4))
    except ValueError:
        return None


def _set_shading(element, color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color.strip().lstrip("#"))
    element.get_or_add_tcPr().append(shd)


def _apply_run_style(run, attrs: dict[str, Any]) -> None:
    if attrs.get("bold"):
        run.bold = True
    if attrs.get("italic"):
        run.italic = True
    color = _safe_color(attrs.get("color"))
    if color:
        run.font.color.rgb = RGBColor(*color)
    size = attrs.get("size")
    if isinstance(size, int) and size > 0:
        run.font.size = size


def _create_paragraph(parent, text: str, attrs: dict[str, Any]) -> None:
    paragraph = parent.add_paragraph()
    if attrs.get("align") in {"center", "right", "left", "justify"}:
        paragraph.alignment = {
            "left": 0,
            "center": 1,
            "right": 2,
            "justify": 3,
        }[attrs["align"]]
    run = paragraph.add_run(text)
    _apply_run_style(run, attrs)


def _create_heading(doc: Document, block: dict[str, Any]) -> None:
    level = block.get("level")
    if not isinstance(level, int) or level < 1 or level > 4:
        level = 1
    text = str(block.get("text", "")).strip()
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    if text:
        run = paragraph.add_run(text)
        _apply_run_style(run, block)
    color = _safe_color(block.get("color"))
    if color:
        paragraph.runs[0].font.color.rgb = RGBColor(*color)


def _create_paragraph_block(doc: Document, block: dict[str, Any]) -> None:
    text = str(block.get("text", "")).strip()
    if not text:
        return
    _create_paragraph(doc, text, block)


def _create_list(doc: Document, block: dict[str, Any], ordered: bool = False) -> None:
    items = block.get("items")
    if not isinstance(items, list):
        return
    bullet_char = str(block.get("bullet", "•")) if not ordered else None
    for item in items:
        if isinstance(item, dict):
            text = str(item.get("text", "")).strip()
            attrs = item
        else:
            text = str(item).strip()
            attrs = {}
        if not text:
            continue
        paragraph = doc.add_paragraph(style="List Paragraph")
        if ordered:
            paragraph.style = "List Number"
        run = paragraph.add_run(text)
        _apply_run_style(run, {**block, **attrs})
        if bullet_char and paragraph.runs:
            paragraph.runs[0].text = f"{bullet_char} {paragraph.runs[0].text}"


def _create_table(doc: Document, block: dict[str, Any]) -> None:
    headers = block.get("headers")
    rows = block.get("rows")
    if not isinstance(headers, list) or not isinstance(rows, list):
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        paragraph = cell.paragraphs[0]
        paragraph.text = str(header)
        color = _safe_color(block.get("header_color"))
        if color:
            paragraph.runs[0].font.color.rgb = RGBColor(*color)
        header_bg = block.get("header_bg")
        if isinstance(header_bg, str):
            _set_shading(cell._tc, header_bg)
    for row_idx, row in enumerate(rows, start=1):
        if not isinstance(row, list):
            continue
        for col_idx, value in enumerate(row[: len(headers)]):
            cell = table.rows[row_idx].cells[col_idx]
            paragraph = cell.paragraphs[0]
            text = str(value)
            paragraph.text = text
            paragraph.runs[0].font.size = None


def _create_block(doc: Document, block: dict[str, Any]) -> None:
    content = block.get("content")
    bg = block.get("bg")
    if isinstance(content, list):
        paragraph = doc.add_paragraph()
        if isinstance(bg, str):
            _set_shading(paragraph._p, bg)
        for child in content:
            if not isinstance(child, dict):
                continue
            _dispatch_block(doc, child)
    elif isinstance(content, str):
        paragraph = doc.add_paragraph(str(content))
        if isinstance(bg, str):
            _set_shading(paragraph._p, bg)


def _create_divider(doc: Document, block: dict[str, Any]) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("\n")
    if color := _safe_color(block.get("color")):
        run.font.color.rgb = RGBColor(*color)


def _create_spacer(doc: Document, block: dict[str, Any]) -> None:
    size = block.get("size")
    if isinstance(size, int) and size > 0:
        paragraph = doc.add_paragraph()
        paragraph.add_run("\n" * (size // 12))


def _create_page_break(doc: Document, block: dict[str, Any]) -> None:
    doc.add_page_break()


def _dispatch_block(doc: Document, block: dict[str, Any]) -> None:
    if not isinstance(block, dict):
        return
    block_type = block.get("type")
    try:
        if block_type == "heading":
            _create_heading(doc, block)
        elif block_type == "paragraph":
            _create_paragraph_block(doc, block)
        elif block_type == "bullet_list":
            _create_list(doc, block, ordered=False)
        elif block_type == "numbered_list":
            _create_list(doc, block, ordered=True)
        elif block_type == "table":
            _create_table(doc, block)
        elif block_type == "block":
            _create_block(doc, block)
        elif block_type == "divider":
            _create_divider(doc, block)
        elif block_type == "spacer":
            _create_spacer(doc, block)
        elif block_type == "page_break":
            _create_page_break(doc, block)
        else:
            logger.warning("Skipping unknown docx block type: %s", block_type)
    except Exception as exc:
        logger.warning("Skipping malformed docx block %s: %s", block_type, exc)


def render_docx(document: dict[str, Any]) -> tuple[bytes, str]:
    doc = Document()
    title = document.get("title")
    if isinstance(title, str) and title.strip():
        doc.add_heading(title.strip(), level=1)
    blocks = document.get("blocks")
    if isinstance(blocks, list):
        for block in blocks:
            _dispatch_block(doc, block)
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
