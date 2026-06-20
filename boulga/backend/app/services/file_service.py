"""
FileService — stockage, téléchargement et préparation des fichiers pour le LLM.

Types de retour de get_file_content_for_llm :
  {"type": "binary", "data": bytes, "mime_type": str, "name": str}  ← multimodal
  {"type": "text",   "content": str, "name": str}                   ← injection texte
"""
import csv
import io
import uuid
from typing import Optional
from uuid import UUID

from app.db.repositories.file_repository import FileRepository
from app.db.session import get_supabase

# ── Constantes ───────────────────────────────────────────────────────────────

STORAGE_BUCKET = "submited-files"
GENERATED_BUCKET = "generated-files"
MAX_TEXT_CHARS = 200_000        # cap sur le texte extrait envoyé au LLM
CSV_LARGE_THRESHOLD = 50_000    # lignes au-delà desquelles on envoie un profil + échantillon
CSV_SAMPLE_ROWS = 20

# Providers supportant les fichiers binaires (multimodal) pour les PDF
MULTIMODAL_PROVIDERS = {"gemini", "openai"}

# Extensions → MIME types normalisés (pour les uploads avec MIME incorrects)
_EXT_MIME: dict[str, str] = {
    ".csv":  "text/csv",
    ".txt":  "text/plain",
    ".pdf":  "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".png":  "image/png",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif":  "image/gif",
    ".webp": "image/webp",
}


def _normalize_mime(filename: str, declared_mime: str) -> str:
    """Retourne le MIME type fiable en combinant extension et type déclaré."""
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return _EXT_MIME.get(ext, declared_mime or "application/octet-stream")


# ── Service ──────────────────────────────────────────────────────────────────

class FileService:
    def __init__(self) -> None:
        self._db = get_supabase()
        self._repo = FileRepository(self._db)

    # ── Storage ───────────────────────────────────────────────────────────

    def _upload(self, path: str, content: bytes, mime_type: str) -> None:
        self._db.storage.from_(STORAGE_BUCKET).upload(
            path=path,
            file=content,
            file_options={"contentType": mime_type},
        )

    def _download(self, storage_path: str) -> bytes:
        return self._db.storage.from_(STORAGE_BUCKET).download(storage_path)

    # ── Public API ────────────────────────────────────────────────────────

    def store_file(
        self,
        user_id: str,
        filename: str,
        content: bytes,
        mime_type: str,
    ) -> dict:
        """
        Stocke le fichier dans Supabase Storage et crée l'enregistrement en base.
        Retourne le dict de la table files (avec id, original_name, mime_type, size_bytes).
        """
        file_id = str(uuid.uuid4())
        mime_type = _normalize_mime(filename, mime_type)
        storage_path = f"{user_id}/{file_id}/{filename}"

        self._upload(storage_path, content, mime_type)

        return self._repo.create({
            "id": file_id,
            "user_id": user_id,
            "original_name": filename,
            "storage_path": storage_path,
            "mime_type": mime_type,
            "size_bytes": len(content),
        })

    def list_for_user(self, user_id: str, limit: int = 50) -> list[dict]:
        """Liste les fichiers d'un utilisateur, triés par date décroissante."""
        return self._repo.list_by_user(UUID(user_id), limit=limit)

    def get_meta(self, file_id: str) -> Optional[dict]:
        """Retourne les métadonnées du fichier depuis la DB."""
        return self._repo.get_by_id(UUID(file_id))

    def download_content(self, storage_path: str) -> bytes:
        """Télécharge le fichier depuis Supabase Storage (upload ou generated)."""
        if storage_path.startswith("generated/"):
            path_in_bucket = storage_path[len("generated/"):]
            return bytes(self._db.storage.from_(GENERATED_BUCKET).download(path_in_bucket))
        return self._download(storage_path)

    def store_generated_file(
        self,
        user_id: str,
        filename: str,
        content: bytes,
        mime_type: str,
        conversation_id: str | None = None,
        message_id: str | None = None,
    ) -> dict:
        """
        Stocke un fichier généré par le LLM dans le bucket Supabase 'generated'.
        Retourne le dict de la table files (id, original_name, mime_type, size_bytes…).
        Les colonnes source/conversation_id/message_id sont omises pour compatibilité
        avec le schéma initial — l'utilisateur les ajoutera manuellement dans Supabase.
        """
        import uuid as _uuid
        file_id = str(_uuid.uuid4())
        path_in_bucket = f"{user_id}/{file_id}/{filename}"
        storage_path = f"generated/{path_in_bucket}"

        self._db.storage.from_(GENERATED_BUCKET).upload(
            path=path_in_bucket,
            file=content,
            file_options={"content-type": mime_type},
        )

        data: dict = {
            "id": file_id,
            "user_id": user_id,
            "original_name": filename,
            "storage_path": storage_path,
            "mime_type": mime_type,
            "size_bytes": len(content),
        }

        return self._repo.create(data)

    def get_signed_url(self, file_id: str, expires_in: int = 3600) -> str | None:
        """Retourne une URL signée Supabase (valide expires_in secondes) pour un fichier généré."""
        try:
            meta = self.get_meta(file_id)
            if not meta:
                return None
            storage_path = meta["storage_path"]
            if not storage_path.startswith("generated/"):
                return None
            path_in_bucket = storage_path[len("generated/"):]
            result = self._db.storage.from_(GENERATED_BUCKET).create_signed_url(
                path=path_in_bucket,
                expires_in=expires_in,
            )
            # Supabase SDK retourne {"signedURL": "https://..."} ou {"signed_url": ...}
            return result.get("signedURL") or result.get("signed_url") or None
        except Exception:
            return None

    def get_file_content_for_llm(
        self,
        file_id: str,
        provider: str = "gemini",
    ) -> dict:
        """
        Retourne le contenu du fichier sous une forme adaptée au LLM.

        Stratégie par type :
          Image       → binary (multimodal) pour tous providers
          PDF         → binary si provider multimodal, sinon extraction texte
          CSV / TXT   → text (avec profil + échantillon si > 50 000 lignes)
          DOCX        → text (paragraphes + tableaux)
          XLSX        → text (CSV par feuille)
          Autre       → binary fallback
        """
        meta = self._repo.get_by_id(UUID(file_id))
        if not meta:
            return {}

        raw = self._download(meta["storage_path"])
        mime = _normalize_mime(meta.get("original_name", ""), meta.get("mime_type", ""))
        name = meta.get("original_name", "fichier")

        # ── Images ─────────────────────────────────────────────────────
        if mime.startswith("image/"):
            return {"type": "binary", "data": raw, "mime_type": mime, "name": name}

        # ── PDF ────────────────────────────────────────────────────────
        if mime == "application/pdf":
            if provider in MULTIMODAL_PROVIDERS:
                return {"type": "binary", "data": raw, "mime_type": mime, "name": name}
            text = self._extract_pdf_text(raw)
            return {"type": "text", "content": text, "name": name}

        # ── CSV ────────────────────────────────────────────────────────
        if mime in ("text/csv", "application/csv") or name.lower().endswith(".csv"):
            text = self._extract_csv(raw, name)
            return {"type": "text", "content": text, "name": name}

        # ── TXT / texte brut ───────────────────────────────────────────
        if mime.startswith("text/"):
            text = self._extract_plain_text(raw)
            return {"type": "text", "content": text, "name": name}

        # ── DOCX ───────────────────────────────────────────────────────
        if (
            mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or name.lower().endswith(".docx")
        ):
            text = self._extract_docx(raw)
            return {"type": "text", "content": text, "name": name}

        # ── XLSX ───────────────────────────────────────────────────────
        if (
            mime in (
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/vnd.ms-excel",
            )
            or name.lower().endswith((".xlsx", ".xls"))
        ):
            text = self._extract_xlsx(raw)
            return {"type": "text", "content": text, "name": name}

        # ── Fallback binaire ───────────────────────────────────────────
        return {"type": "binary", "data": raw, "mime_type": mime, "name": name}

    def get_file_profile(self, file_id: str) -> dict:
        """
        Retourne un profil résumé du fichier.
        Pour les CSV : nb lignes, colonnes, valeurs manquantes.
        """
        meta = self._repo.get_by_id(UUID(file_id))
        if not meta:
            return {}

        name = meta.get("original_name", "")
        mime = _normalize_mime(name, meta.get("mime_type", ""))

        profile: dict = {
            "id": file_id,
            "name": name,
            "mime_type": mime,
            "size_bytes": meta.get("size_bytes", 0),
        }

        if mime in ("text/csv", "application/csv") or name.lower().endswith(".csv"):
            try:
                raw = self._download(meta["storage_path"])
                text = raw.decode("utf-8", errors="replace")
                rows = list(csv.reader(io.StringIO(text)))
                headers = rows[0] if rows else []
                data_rows = rows[1:]

                null_counts = {
                    col: sum(
                        1 for row in data_rows
                        if idx >= len(row) or row[idx].strip() == ""
                    )
                    for idx, col in enumerate(headers)
                }

                profile.update({
                    "total_rows": len(data_rows),
                    "columns": len(headers),
                    "column_names": headers,
                    "null_counts": null_counts,
                })
            except Exception as exc:
                profile["error"] = str(exc)

        return profile

    # ── Extracteurs ───────────────────────────────────────────────────────

    def _extract_csv(self, content: bytes, name: str) -> str:
        """
        Extrait le contenu CSV.
        Si > 50 000 lignes : retourne le profil + les 20 premières lignes de données.
        """
        try:
            text = content.decode("utf-8", errors="replace")
        except Exception:
            return f"[Impossible de décoder le fichier {name}]"

        try:
            reader = csv.reader(io.StringIO(text))
            rows = list(reader)
        except Exception as exc:
            return f"[Erreur lecture CSV : {exc}]"

        if not rows:
            return f"[Fichier CSV vide : {name}]"

        headers = rows[0]
        data_rows = rows[1:]
        total = len(data_rows)

        # Fichier petit → contenu complet
        if total <= CSV_LARGE_THRESHOLD:
            return text[:MAX_TEXT_CHARS]

        # Grand fichier → profil + échantillon
        sample = data_rows[:CSV_SAMPLE_ROWS]

        null_counts = {
            col: sum(
                1 for row in data_rows
                if idx >= len(row) or row[idx].strip() == ""
            )
            for idx, col in enumerate(headers)
        }

        out = io.StringIO()
        writer = csv.writer(out)
        writer.writerow(headers)
        writer.writerows(sample)
        sample_csv = out.getvalue()

        null_summary = " | ".join(
            f"{col}: {n}" for col, n in null_counts.items() if n > 0
        ) or "aucune"

        return (
            f"[Fichier CSV volumineux : {name}]\n"
            f"Lignes de données : {total:,}\n"
            f"Colonnes ({len(headers)}) : {', '.join(headers)}\n"
            f"Valeurs manquantes : {null_summary}\n\n"
            f"--- Échantillon ({CSV_SAMPLE_ROWS} premières lignes) ---\n"
            f"{sample_csv}"
        )

    def _extract_plain_text(self, content: bytes) -> str:
        """Extrait le texte brut depuis des bytes."""
        try:
            text = content.decode("utf-8", errors="replace")
        except Exception:
            return "[Impossible de décoder le fichier texte]"

        lines = text.splitlines()
        if len(lines) <= CSV_LARGE_THRESHOLD:
            return text[:MAX_TEXT_CHARS]

        # Très long fichier texte : troncature avec avertissement
        sample = "\n".join(lines[:1000])
        return (
            f"[Fichier texte volumineux — {len(lines):,} lignes]\n"
            f"--- Extrait des 1 000 premières lignes ---\n"
            f"{sample}"
        )[:MAX_TEXT_CHARS]

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extrait le texte d'un PDF via pdfplumber (fallback pour providers non-multimodaux)."""
        try:
            import pdfplumber
        except ImportError:
            return "[Extraction PDF indisponible — installez pdfplumber]"

        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages_text = []
                for i, page in enumerate(pdf.pages):
                    t = page.extract_text()
                    if t:
                        pages_text.append(f"[Page {i + 1}]\n{t}")
            return ("\n\n".join(pages_text))[:MAX_TEXT_CHARS]
        except Exception as exc:
            return f"[Erreur extraction PDF : {exc}]"

    def _extract_docx(self, content: bytes) -> str:
        """Extrait paragraphes et tableaux d'un fichier DOCX."""
        try:
            from docx import Document
        except ImportError:
            return "[Extraction DOCX indisponible — installez python-docx]"

        try:
            doc = Document(io.BytesIO(content))
            parts: list[str] = []

            for para in doc.paragraphs:
                stripped = para.text.strip()
                if stripped:
                    parts.append(stripped)

            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(cells))
                if table_rows:
                    parts.append("\n".join(table_rows))

            return ("\n\n".join(parts))[:MAX_TEXT_CHARS]
        except Exception as exc:
            return f"[Erreur extraction DOCX : {exc}]"

    def _extract_xlsx(self, content: bytes) -> str:
        """Extrait les feuilles d'un fichier XLSX en format CSV par feuille."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            return "[Extraction XLSX indisponible — installez openpyxl]"

        try:
            wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            sheets: list[str] = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                out = io.StringIO()
                writer = csv.writer(out)
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    writer.writerow(
                        ["" if v is None else str(v) for v in row]
                    )
                    row_count += 1
                    if row_count >= CSV_LARGE_THRESHOLD:
                        break

                sheet_csv = out.getvalue()
                if sheet_csv.strip():
                    header = f"# Feuille : {sheet_name}"
                    if row_count >= CSV_LARGE_THRESHOLD:
                        header += f" (tronqué à {CSV_LARGE_THRESHOLD:,} lignes)"
                    sheets.append(f"{header}\n{sheet_csv}")

            wb.close()
            return ("\n\n".join(sheets))[:MAX_TEXT_CHARS]
        except Exception as exc:
            return f"[Erreur extraction XLSX : {exc}]"
