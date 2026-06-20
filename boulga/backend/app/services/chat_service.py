import re as _re
from datetime import date
from typing import AsyncIterator, Optional
from uuid import UUID

from app.db.repositories.conversation_repository import ConversationRepository
from app.db.repositories.message_repository import MessageRepository
from app.db.session import get_supabase
from app.manager.llm_manager import llm_manager
from app.manager.router_agent import router_agent
from app.prompts.chat_prompts import FILE_GENERATION_ADDENDUM, IMAGE_GENERATION_ADDENDUM, TITLE_GENERATION_PROMPT
from app.services.image_service import ImageGenerationError, generate_image
from app.prompts.tool_prompts import get_full_system_prompt
from app.services.code_executor import CodeExecutionError, CodeExecutor
from app.services.file_service import FileService
from app.services.quota_service import QuotaService
from app.services.subscription_service import SubscriptionService

# Tiers autorisés à utiliser le Routage Automatique (Source+)
_ROUTING_TIERS: set[str] = {"source", "fleuve", "ocean"}

# ── Constantes ───────────────────────────────────────────────────────────────

# Fenêtres de contexte approximatives par modèle (en tokens)
MODEL_CONTEXT_WINDOWS: dict[str, int] = {
    "gemini-2.5-flash": 1_000_000,
    "gemini-2.5-pro":   2_000_000,
    "claude-haiku-4-5": 200_000,
    "claude-sonnet-4-6":200_000,
    "gpt-5.5-instant":  128_000,
    "gpt-5.5-pro":      200_000,
    "deepseek-v4-flash":128_000,
    "deepseek-v4-pro":  128_000,
}

_ECO_PROVIDER = "gemini"
_ECO_MODEL    = "gemini-2.5-flash"
_RECENT_KEEP  = 10  # messages récents conservés intacts lors d'un résumé de contexte

def _estimate_tokens(text: str) -> int:
    return max(1, len(text) // 4)


# ── Détection d'intention fichier ────────────────────────────────────────────

_FILE_INTENT_RE = _re.compile(
    r"\b(génère|générer|crée|créer|produis|produire|exporte|exporter|"
    r"word|excel|powerpoint|pdf|docx|xlsx|pptx|"
    r"fichier|document|tableau|rapport)\b",
    _re.IGNORECASE | _re.UNICODE,
)

_IMAGE_INTENT_RE = _re.compile(
    r"\b(génère|générer|crée|créer|dessine|dessiner|illustre|illustrer|"
    r"image|photo|illustration|logo|icône|icone|avatar|bannière|banniere|"
    r"affiche|poster|visuel|artwork|portrait|paysage|fond.d.écran)\b",
    _re.IGNORECASE | _re.UNICODE,
)

# Mots qui excluent l'intention image (fichiers bureautiques)
_IMAGE_EXCLUSION_RE = _re.compile(
    r"\b(word|excel|pdf|docx|xlsx|pptx|powerpoint|tableau.excel|feuille|"
    r"fichier|document|rapport|facture|contrat|bilan)\b",
    _re.IGNORECASE | _re.UNICODE,
)


def _wants_file(message: str) -> bool:
    return bool(_FILE_INTENT_RE.search(message))


def _wants_image(message: str) -> bool:
    """Détecte si l'utilisateur veut une image générée (prioritaire sur fichier)."""
    if not _IMAGE_INTENT_RE.search(message):
        return False
    # Si le message mentionne explicitement un format bureautique → pas une image
    if _IMAGE_EXCLUSION_RE.search(message):
        return False
    return True


def _extract_python_blocks(text: str) -> list[str]:
    """Extrait les blocs ```python ... ``` d'un texte (fallback si code_execution inactif)."""
    return _re.findall(r"```python\s*\n(.*?)\n```", text, _re.DOTALL)


# ── Helpers ──────────────────────────────────────────────────────────────────


def _create_file_from_claude_tool(tool_input: dict) -> tuple[bytes, str, str]:
    """Crée un fichier binaire depuis les paramètres de l'outil create_file de Claude."""
    import io as _io

    fmt = tool_input.get("format", "txt")
    content = tool_input.get("content", "")
    filename = tool_input.get("filename", f"document.{fmt}")

    if fmt == "docx":
        from docx import Document
        doc = Document()
        if isinstance(content, str):
            for line in content.split("\n"):
                stripped = line.rstrip()
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped:
                    doc.add_paragraph(stripped)
        buf = _io.BytesIO()
        doc.save(buf)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return buf.getvalue(), mime, filename

    if fmt == "xlsx":
        from openpyxl import Workbook
        wb = Workbook()
        wb.remove(wb.active)
        sheets_data = content if isinstance(content, list) else [{"name": "Feuille1", "rows": []}]
        for sheet_def in sheets_data:
            ws = wb.create_sheet(title=str(sheet_def.get("name", "Feuille"))[:31])
            for row in sheet_def.get("rows", []):
                ws.append(row if isinstance(row, list) else [str(row)])
        buf = _io.BytesIO()
        wb.save(buf)
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        return buf.getvalue(), mime, filename

    # csv / txt / md → texte brut
    text = content if isinstance(content, str) else str(content)
    mime_map = {"csv": "text/csv", "md": "text/markdown", "txt": "text/plain"}
    mime = mime_map.get(fmt, "text/plain")
    return text.encode("utf-8"), mime, filename


# ── Service ──────────────────────────────────────────────────────────────────

class ChatService:
    def __init__(self) -> None:
        db = get_supabase()
        self._conv_repo  = ConversationRepository(db)
        self._msg_repo   = MessageRepository(db)
        self._quota_svc  = QuotaService()
        self._sub_svc    = SubscriptionService()

    # ── Fichiers ───────────────────────────────────────────────────────

    def _prepare_files(
        self,
        file_ids: list[str],
        provider: str,
    ) -> tuple[str, list[dict]]:
        """
        Résout les file_ids en contenu LLM.

        Retourne :
          text_context  — texte extrait à injecter dans le message (CSV, DOCX, XLSX, TXT…)
          binary_files  — fichiers binaires pour envoi multimodal (images, PDF Gemini)
                          format : [{"data": bytes, "mime_type": str, "name": str}]
        """
        text_parts: list[str] = []
        binary_files: list[dict] = []

        if not file_ids:
            return "", []

        svc = FileService()
        for fid in file_ids:
            try:
                fc = svc.get_file_content_for_llm(fid, provider)
                if not fc:
                    continue
                if fc["type"] == "binary":
                    binary_files.append({
                        "data":      fc["data"],
                        "mime_type": fc["mime_type"],
                        "name":      fc.get("name", ""),
                    })
                else:
                    name    = fc.get("name", "fichier")
                    content = fc.get("content", "")
                    text_parts.append(f"[Contenu du fichier : {name}]\n{content}")
            except Exception:
                pass  # Fichier ignoré — le chat continue

        return "\n\n".join(text_parts), binary_files

    # ── Gestion contexte long ──────────────────────────────────────────

    async def _summarize_messages(self, messages: list[dict]) -> str:
        text = "\n".join(
            f"{m.get('role', 'user').upper()}: {m.get('content', '')[:500]}"
            for m in messages
        )
        prompt = (
            "Résume brièvement cette conversation en 4-6 phrases clés "
            "(points importants, décisions prises, contexte nécessaire) :\n\n" + text
        )
        try:
            return await llm_manager.generate_text(
                provider=_ECO_PROVIDER,
                model_id=_ECO_MODEL,
                prompt=prompt,
            )
        except Exception:
            return "[Résumé de l'historique indisponible]"

    async def _prepare_history(
        self, history: list[dict], model_id: str
    ) -> list[dict]:
        """
        Prépare l'historique pour le LLM.
        Si l'historique dépasse 80 % de la fenêtre du modèle, résume les anciens messages
        et conserve les _RECENT_KEEP derniers intacts.
        """
        context_window = MODEL_CONTEXT_WINDOWS.get(model_id, 128_000)
        threshold      = int(context_window * 0.8)
        total_tokens   = sum(_estimate_tokens(m.get("content", "")) for m in history)

        if total_tokens <= threshold or len(history) <= _RECENT_KEEP:
            return history

        recent = history[-_RECENT_KEEP:]
        older  = history[:-_RECENT_KEEP]
        summary = await self._summarize_messages(older)

        return [
            {"role": "user",  "content": f"[Résumé de l'historique précédent : {summary}]"},
            {"role": "model", "content": "Compris, je tiens compte de cet historique."},
        ] + recent

    # ── Streaming principal ────────────────────────────────────────────

    async def stream_message(
        self,
        user_id: str,
        message: str,
        provider: str = "gemini",
        model_id: str = "gemini-2.5-flash",
        conversation_id: Optional[str] = None,
        file_ids: Optional[list[str]] = None,
        tool_slug: Optional[str] = None,
        auto_route: bool = False,
    ) -> AsyncIterator[dict]:
        """
        Async generator qui yield des dicts (événements SSE).

          {"type": "conversation", "id": str, "is_new": bool}
          {"type": "error",        "message": str}            ← stoppe
          {"type": "chunk",        "text": str}               ← 0..N
          {"type": "title",        "title": str}              ← si nouvelle conv.
          {"type": "done",         "message_id": str}
        """
        file_ids = file_ids or []
        is_new   = conversation_id is None

        # a. Créer ou vérifier la conversation
        if is_new:
            conv = self._conv_repo.create({
                "user_id":  user_id,
                "provider": provider,
                "model_id": model_id,
            })
            conversation_id = conv["id"]
        else:
            conv = self._conv_repo.get_by_id(UUID(conversation_id))
            if not conv or conv.get("user_id") != user_id:
                yield {"type": "error", "message": "Conversation introuvable"}
                return

        # b. Événement conversation
        yield {"type": "conversation", "id": conversation_id, "is_new": is_new}

        # c. Vérifications quota + accès modèle (AVANT l'appel LLM)
        if not await self._quota_svc.is_message_allowed(user_id):
            yield {
                "type":    "error",
                "message": "quota_exceeded",
            }
            return

        if not self._sub_svc.check_can_use_model(user_id, provider, model_id):
            yield {
                "type":    "error",
                "message": "model_access_denied",
            }
            return

        # c-bis. Vérification quota image (coupe tôt, avant même le stream texte)
        if _wants_image(message):
            if not await self._quota_svc.is_image_allowed(user_id):
                yield {
                    "type":    "error",
                    "message": "image_quota_exceeded",
                }
                return

        # d. Routage Automatique (si activé et tier Source+)
        tier = self._sub_svc.get_tier(user_id)
        if auto_route and tier in _ROUTING_TIERS:
            try:
                route = await router_agent.route(message, tier)
                provider  = route["provider"]
                model_id  = route["model_id"]
                yield {
                    "type":     "routing",
                    "provider": provider,
                    "model":    model_id,
                    "reason":   route["reason"],
                }
            except Exception:
                pass  # On garde le provider/modèle initial si le routage échoue

        # e. Enregistrer le message utilisateur
        self._msg_repo.create({
            "conversation_id": conversation_id,
            "role":            "user",
            "content":         message,
            "provider":        provider,
            "model_id":        model_id,
        })

        # e. Récupérer l'historique complet
        history = self._msg_repo.list_by_conversation(UUID(conversation_id))

        # f. Préparer les fichiers (texte extrait + binaires multimodaux)
        text_context, binary_files = self._prepare_files(file_ids, provider)

        # g. Assembler le system prompt
        system_prompt = get_full_system_prompt(tool_slug)

        # h. Gestion du contexte long
        history_for_llm = await self._prepare_history(history, model_id)

        # Convertir les rôles DB → rôles LLM ("assistant" → "model" pour Gemini)
        llm_messages = [
            {
                "role": "model" if m.get("role") == "assistant" else m.get("role", "user"),
                "content": m.get("content", ""),
            }
            for m in history_for_llm
        ]

        # Injecter le contenu textuel des fichiers dans le dernier message utilisateur
        if text_context:
            for i in range(len(llm_messages) - 1, -1, -1):
                if llm_messages[i]["role"] == "user":
                    llm_messages[i] = {
                        **llm_messages[i],
                        "content": llm_messages[i]["content"] + f"\n\n{text_context}",
                    }
                    break

        # i. Détecter l'intention (image prioritaire sur fichier bureautique)
        wants_img = _wants_image(message)
        enable_code_exec = _wants_file(message) and not wants_img
        if wants_img:
            system_prompt = system_prompt + IMAGE_GENERATION_ADDENDUM
        elif enable_code_exec:
            system_prompt = system_prompt + FILE_GENERATION_ADDENDUM

        # j. Streamer via llm_manager
        full_response = ""
        accumulated_code: list[str] = []
        gemini_binary_files: list[dict] = []
        claude_tool_calls: list[dict] = []
        # Quand enable_code_exec est actif, masquer les blocs de code Python (UX propre)
        _code_block_depth = 0
        try:
            async for event in llm_manager.stream_chat(
                provider=provider,
                model_id=model_id,
                messages=llm_messages,
                system_prompt=system_prompt,
                files=binary_files if binary_files else None,
                enable_code_execution=enable_code_exec and provider in ("gemini", "claude"),
            ):
                if event["type"] == "text":
                    full_response += event["text"]
                    if enable_code_exec:
                        # Compter les ``` pour savoir si on est dans un bloc de code
                        chunk = event["text"]
                        for line in chunk.splitlines(keepends=True):
                            stripped = line.strip()
                            if stripped.startswith("```"):
                                if _code_block_depth == 0:
                                    _code_block_depth = 1  # entrée dans le bloc
                                else:
                                    _code_block_depth = 0  # sortie du bloc
                            elif _code_block_depth == 0:
                                yield {"type": "chunk", "text": line}
                    else:
                        yield {"type": "chunk", "text": event["text"]}
                elif event["type"] == "code":
                    accumulated_code.append(event["code"])
                    # Code natif Gemini/Claude : ne pas afficher le code technique
                elif event["type"] == "code_result":
                    if event.get("output"):
                        yield {"type": "chunk", "text": f"\n*Résultat :* {event['output']}\n"}
                elif event["type"] == "file_data":
                    # Gemini a généré un fichier binaire directement via code_execution
                    from datetime import timezone
                    ext_map = {
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
                        "application/pdf": "pdf",
                        "text/csv": "csv",
                        "text/plain": "txt",
                    }
                    mime_type = event["mime_type"]
                    ext = ext_map.get(mime_type, "bin")
                    filename = f"document_{date.today().strftime('%Y%m%d')}.{ext}"
                    gemini_binary_files.append({
                        "data": event["data"],
                        "mime_type": mime_type,
                        "filename": filename,
                    })
                elif event["type"] == "tool_create_file":
                    claude_tool_calls.append(event["tool_input"])
        except Exception as exc:
            yield {"type": "error", "message": getattr(exc, "message", str(exc))}
            return

        # k. Enregistrer la réponse de l'assistant
        assistant_msg = self._msg_repo.create({
            "conversation_id": conversation_id,
            "role":            "assistant",
            "content":         full_response,
            "provider":        provider,
            "model_id":        model_id,
        })

        # l. Consommer le quota (message + tokens estimés)
        try:
            await self._quota_svc.consume_message(user_id)
            tokens_est = _estimate_tokens(full_response)
            await self._quota_svc.consume_tokens(user_id, tokens_est)
        except Exception:
            pass

        # l-image. Générer l'image si demandée
        if wants_img:
            try:
                await self._quota_svc.consume_image(user_id)
                img_result = await generate_image(
                    provider=provider,
                    prompt=message,
                    user_id=user_id,
                    conversation_id=conversation_id,
                    message_id=assistant_msg["id"] if assistant_msg else None,
                )
                # Note discrète si fallback vers GPT Image
                if img_result["used_fallback"]:
                    yield {
                        "type": "chunk",
                        "text": "\n\n*Image générée via GPT Image.*",
                    }
                yield {
                    "type":       "file_ready",
                    "file_id":    img_result["file_id"],
                    "filename":   img_result["filename"],
                    "mime_type":  "image/png",
                    "size_bytes": img_result["size_bytes"],
                    "url":        img_result["url"],
                    "is_image":   True,
                }
            except ImageGenerationError as img_err:
                yield {"type": "error", "message": f"Génération image échouée : {img_err}"}
            except Exception as img_err:
                yield {"type": "error", "message": f"Erreur image : {img_err}"}

        # l-bis. Stocker les fichiers binaires Gemini (code_execution inline_data)
        if gemini_binary_files:
            file_svc = FileService()
            for gf in gemini_binary_files:
                try:
                    record = file_svc.store_generated_file(
                        user_id=user_id,
                        filename=gf["filename"],
                        content=gf["data"],
                        mime_type=gf["mime_type"],
                        conversation_id=conversation_id,
                        message_id=assistant_msg["id"] if assistant_msg else None,
                    )
                    yield {
                        "type": "file_ready",
                        "file_id": record["id"],
                        "filename": gf["filename"],
                        "mime_type": gf["mime_type"],
                        "size_bytes": len(gf["data"]),
                        "url": f"/api/files/{record['id']}/download",
                    }
                except Exception as store_err:
                    yield {"type": "error", "message": f"Erreur stockage fichier Gemini : {store_err}"}

        # l-ter. Créer les fichiers depuis les appels d'outil Claude (create_file)
        if claude_tool_calls:
            file_svc = FileService()
            for tool_input in claude_tool_calls:
                try:
                    file_bytes, mime_type, filename = _create_file_from_claude_tool(tool_input)
                    record = file_svc.store_generated_file(
                        user_id=user_id,
                        filename=filename,
                        content=file_bytes,
                        mime_type=mime_type,
                        conversation_id=conversation_id,
                        message_id=assistant_msg["id"] if assistant_msg else None,
                    )
                    yield {
                        "type": "file_ready",
                        "file_id": record["id"],
                        "filename": filename,
                        "mime_type": mime_type,
                        "size_bytes": len(file_bytes),
                        "url": f"/api/files/{record['id']}/download",
                    }
                except Exception as e:
                    yield {"type": "error", "message": f"Erreur création fichier Claude : {e}"}

        # l-quater. Exécuter le code généré et stocker les fichiers produits
        if enable_code_exec and not accumulated_code:
            accumulated_code = _extract_python_blocks(full_response)

        if accumulated_code:
            executor = CodeExecutor()
            file_svc = FileService()
            for code_block in accumulated_code:
                try:
                    produced = await executor.execute(code_block, timeout=30.0)
                except CodeExecutionError as exec_err:
                    yield {"type": "error", "message": f"Exécution échouée : {exec_err}"}
                    continue
                for file_info in produced:
                    try:
                        record = file_svc.store_generated_file(
                            user_id=user_id,
                            filename=file_info["filename"],
                            content=file_info["content"],
                            mime_type=file_info["mime_type"],
                            conversation_id=conversation_id,
                            message_id=assistant_msg["id"] if assistant_msg else None,
                        )
                        yield {
                            "type": "file_ready",
                            "file_id": record["id"],
                            "filename": file_info["filename"],
                            "mime_type": file_info["mime_type"],
                            "size_bytes": len(file_info["content"]),
                            "url": f"/api/files/{record['id']}/download",
                        }
                    except Exception as store_err:
                        yield {"type": "error", "message": f"Erreur stockage fichier : {store_err}"}

        # m. Générer le titre si nouvelle conversation
        if is_new and full_response:
            try:
                title_prompt = TITLE_GENERATION_PROMPT.format(message=message[:300])
                raw_title = await llm_manager.generate_text(
                    provider=_ECO_PROVIDER,
                    model_id=_ECO_MODEL,
                    prompt=title_prompt,
                )
                title = raw_title.strip().strip('"').strip("'")[:100]
                self._conv_repo.update_title(UUID(conversation_id), title)
                yield {"type": "title", "title": title}
            except Exception:
                pass

        # n. Done
        yield {"type": "done", "message_id": assistant_msg["id"]}
