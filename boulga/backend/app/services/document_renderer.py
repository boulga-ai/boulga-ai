"""document_renderer.py — Rendu de blocs JSON en DOCX/PDF avec 5 templates professionnels."""

import io
from copy import deepcopy
from dataclasses import dataclass
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, KeepTogether,
)
from reportlab.lib.colors import HexColor, grey

from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_XLSX_MIME = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

BULLET_SYMBOLS: dict[str, str] = {
    "dot":    "•",
    "star":   "★",
    "square": "■",
    "check":  "✓",
    "arrow":  "→",
}

# Styles sémantiques pour les callouts — couleurs indépendantes du template
CALLOUT_TYPES: dict[str, dict] = {
    "info":    {"bg": "#E3F2FD", "border": "#1565C0", "label_default": "Information"},
    "tip":     {"bg": "#E8F5E9", "border": "#2E7D32", "label_default": "Conseil"},
    "warning": {"bg": "#FFF8E1", "border": "#F57C00", "label_default": "Attention"},
    "danger":  {"bg": "#FFEBEE", "border": "#C62828", "label_default": "Important"},
    "success": {"bg": "#F1F8E9", "border": "#388E3C", "label_default": "Resultat"},
    "note":    {"bg": "#F5F5F5", "border": "#616161", "label_default": "Note"},
}


# ─── Template definitions ──────────────────────────────────────────────────────

@dataclass
class TemplateColors:
    primary:    str
    secondary:  str
    accent:     str
    light:      str
    text_dark:  str = "#1A1A1A"
    text_light: str = "#FFFFFF"


@dataclass
class TemplateConfig:
    name:             str
    colors:           TemplateColors
    has_page_header:  bool
    bullet_default:   str
    page_margin_cm:   tuple  # (top, bottom, left, right)


TEMPLATES: dict[str, TemplateConfig] = {
    "commercial": TemplateConfig(
        name="commercial",
        colors=TemplateColors(
            primary="#0B1F3A", secondary="#1565C0",
            accent="#F57C00",  light="#E8EEF7",
        ),
        has_page_header=True,
        bullet_default="check",
        page_margin_cm=(2.5, 2.0, 2.5, 2.0),
    ),
    "rapport": TemplateConfig(
        name="rapport",
        colors=TemplateColors(
            primary="#37474F", secondary="#1976D2",
            accent="#2E7D32",  light="#ECEFF1",
        ),
        has_page_header=True,
        bullet_default="square",
        page_margin_cm=(2.5, 2.0, 2.5, 2.0),
    ),
    "contrat": TemplateConfig(
        name="contrat",
        colors=TemplateColors(
            primary="#1A1A1A", secondary="#B71C1C",
            accent="#757575",  light="#F8F8F8",
        ),
        has_page_header=True,
        bullet_default="dot",
        page_margin_cm=(3.0, 2.5, 3.0, 2.5),
    ),
    "rh": TemplateConfig(
        name="rh",
        colors=TemplateColors(
            primary="#1B5E20", secondary="#00695C",
            accent="#F9A825",  light="#F1F8E9",
        ),
        has_page_header=True,
        bullet_default="check",
        page_margin_cm=(2.5, 2.0, 2.5, 2.0),
    ),
    "minimal": TemplateConfig(
        name="minimal",
        colors=TemplateColors(
            primary="#212121", secondary="#1565C0",
            accent="#757575",  light="#F5F5F5",
        ),
        has_page_header=False,
        bullet_default="dot",
        page_margin_cm=(2.0, 2.0, 2.5, 2.5),
    ),
}


# ─── Public API ────────────────────────────────────────────────────────────────

def render_document(
    blocks: list[dict],
    fmt: str,
    template: str = "minimal",
    options: Optional[dict] = None,
) -> tuple[bytes, str]:
    cfg = deepcopy(TEMPLATES.get(template, TEMPLATES["minimal"]))
    opts = options or {}

    if opts.get("primary_color"):
        cfg.colors.primary = opts["primary_color"]

    company_name: str = opts.get("company_name", "") or ""

    if fmt == "docx":
        return _render_docx(blocks, cfg, company_name), _DOCX_MIME
    if fmt == "pdf":
        return _render_pdf(blocks, cfg, company_name), "application/pdf"
    if fmt == "xlsx":
        return _render_xlsx(blocks, cfg, company_name), _XLSX_MIME
    raise ValueError(f"Format non supporté : {fmt}")


# ─── Shared helpers ────────────────────────────────────────────────────────────

def _hex_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _rgb_color(hex_color: str) -> RGBColor:
    r, g, b = _hex_rgb(hex_color)
    return RGBColor(r, g, b)


def _hc(hex_color: str) -> HexColor:
    return HexColor(hex_color)


# ─── DOCX low-level helpers ────────────────────────────────────────────────────

def _para_bg(para, hex_color: str) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    pPr.append(shd)


def _cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _para_left_border(para, hex_color: str, sz: int = 18) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color.lstrip("#"))
    pBdr.append(left)
    pPr.append(pBdr)


def _para_bottom_border(para, hex_color: str, sz: int = 6) -> None:
    """Ligne fine sous un paragraphe — utilisé sous les H1 du template rapport."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _docx_color_band(
    doc: Document,
    text: str,
    bg: str,
    fg: str,
    size: int = 11,
    bold: bool = True,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    pad_v: int = 10,
) -> None:
    """Bande colorée pleine largeur (table 1×1 sans bordures)."""
    avail = (
        doc.sections[0].page_width
        - doc.sections[0].left_margin
        - doc.sections[0].right_margin
    )
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = avail
    cell = tbl.cell(0, 0)
    _cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(pad_v)
    p.paragraph_format.space_after  = Pt(pad_v)
    if text:
        r = p.add_run(text)
        r.font.size  = Pt(size)
        r.font.bold  = bold
        r.font.color.rgb = _rgb_color(fg)


def _docx_spacer(doc: Document, n: int = 1) -> None:
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)


# ─── DOCX page header (en-tête courant) ───────────────────────────────────────

def _docx_page_header(doc: Document, cfg: TemplateConfig, company_name: str) -> None:
    if not cfg.has_page_header and not company_name:
        return
    section = doc.sections[0]
    header  = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    if cfg.name == "commercial":
        _para_bg(p, cfg.colors.primary)
        run = p.add_run(company_name or "Boulga AI")
        run.font.color.rgb = _rgb_color(cfg.colors.text_light)
        run.font.bold = True
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)

    elif cfg.name == "rh":
        _para_bg(p, cfg.colors.primary)
        label = f"CONFIDENTIEL — {company_name}" if company_name else "Ressources Humaines"
        run = p.add_run(label)
        run.font.color.rgb = _rgb_color(cfg.colors.text_light)
        run.font.bold = True
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    elif cfg.name == "rapport":
        _para_bg(p, cfg.colors.light)
        run = p.add_run(company_name or "Rapport")
        run.font.color.rgb = _rgb_color(cfg.colors.primary)
        run.font.italic = True
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Pt(6)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

    elif cfg.name == "contrat":
        run = p.add_run(company_name or "Document confidentiel")
        run.font.color.rgb = _rgb_color(cfg.colors.accent)
        run.font.italic = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    else:  # minimal
        if company_name:
            run = p.add_run(company_name)
            run.font.color.rgb = _rgb_color(cfg.colors.accent)
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ─── DOCX cover page (page de garde) ─────────────────────────────────────────

def _docx_cover_page(
    doc: Document, block: dict, cfg: TemplateConfig, company_name: str
) -> None:
    """Génère la page de garde et ajoute un saut de page à la fin."""
    # La 1re page n'affiche pas l'en-tête courant
    doc.sections[0].different_first_page_header_footer = True

    title       = block.get("title", "Document")
    subtitle    = block.get("subtitle", "")
    author      = block.get("author", "")
    institution = block.get("institution", "") or company_name
    date        = block.get("date", "")
    ref         = block.get("reference", "")
    doc_type    = block.get("doc_type", "")
    meta        = [x for x in [author, date, ref] if x]

    def cp(
        txt: str, size: int,
        bold: bool = False, italic: bool = False,
        color: str = "",
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
        sb: int = 4, sa: int = 4,
    ) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        r = p.add_run(txt)
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        r.font.color.rgb = _rgb_color(color or cfg.colors.primary)

    if cfg.name == "commercial":
        head = "  —  ".join(filter(None, [institution or "Document", doc_type]))
        _docx_color_band(doc, head or "Document Commercial",
                         cfg.colors.primary, cfg.colors.text_light,
                         size=11, bold=True, pad_v=12)
        _docx_spacer(doc, 6)
        cp(title, 26, bold=True, color=cfg.colors.primary, sa=10)
        if subtitle:
            cp(subtitle, 14, italic=True, color=cfg.colors.secondary, sa=4)
        _docx_spacer(doc, 6)
        _docx_color_band(doc, "", cfg.colors.secondary, cfg.colors.text_light, pad_v=2)
        if meta:
            cp("  ·  ".join(meta), 10, color=cfg.colors.accent, sb=8)

    elif cfg.name == "rapport":
        if institution:
            _docx_color_band(doc, institution, cfg.colors.light, cfg.colors.primary,
                             size=11, bold=False, pad_v=8)
        _docx_color_band(doc, "", cfg.colors.secondary, cfg.colors.text_light, pad_v=2)
        _docx_spacer(doc, 6)
        cp(title, 24, bold=True, color=cfg.colors.primary, sa=10)
        if subtitle:
            cp(subtitle, 13, italic=True, color=cfg.colors.secondary, sa=4)
        _docx_spacer(doc, 5)
        _docx_color_band(doc, "", cfg.colors.secondary, cfg.colors.text_light, pad_v=2)
        _docx_spacer(doc, 1)
        if meta:
            cp("  ·  ".join(meta), 10, color=cfg.colors.accent, sb=4)

    elif cfg.name == "contrat":
        _docx_spacer(doc, 4)
        if doc_type:
            cp(doc_type.upper(), 11, bold=True, color=cfg.colors.accent, sa=6)
        _docx_spacer(doc, 2)
        cp(title, 22, bold=True, color=cfg.colors.primary, sa=8)
        if subtitle:
            cp(subtitle, 12, italic=True, color=cfg.colors.accent, sa=4)
        _docx_spacer(doc, 1)
        _docx_color_band(doc, "", cfg.colors.secondary, cfg.colors.text_light, pad_v=2)
        _docx_spacer(doc, 2)
        if institution:
            cp(institution, 11, bold=True, color=cfg.colors.primary, sa=4)
        _docx_spacer(doc, 6)
        if meta:
            cp("  |  ".join(meta), 10, color=cfg.colors.accent, sb=4)

    elif cfg.name == "rh":
        head = "  —  ".join(filter(None, [institution or "Ressources Humaines", doc_type]))
        _docx_color_band(doc, head, cfg.colors.primary, cfg.colors.text_light,
                         size=10, bold=True, pad_v=10)
        _docx_spacer(doc, 5)
        cp(title, 22, bold=True, color=cfg.colors.primary, sa=8)
        if subtitle:
            cp(subtitle, 13, italic=True, color=cfg.colors.secondary, sa=4)
        _docx_spacer(doc, 4)
        if author:
            cp(author, 11, bold=True, color=cfg.colors.primary, sa=4)
        meta2 = [x for x in [date, ref] if x]
        if meta2:
            cp("  ·  ".join(meta2), 10, color=cfg.colors.accent, sb=4)

    else:  # minimal
        _docx_spacer(doc, 8)
        cp(title, 28, bold=True, color=cfg.colors.primary, sa=10)
        if subtitle:
            cp(subtitle, 14, italic=True, color=cfg.colors.secondary, sa=4)
        _docx_spacer(doc, 1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("─" * 40)
        r.font.size = Pt(10)
        r.font.color.rgb = _rgb_color(cfg.colors.accent)
        p.paragraph_format.space_after = Pt(6)
        if meta:
            cp("  ·  ".join(meta), 10, color=cfg.colors.accent, sb=4)

    doc.add_page_break()


# ─── DOCX block renderers ─────────────────────────────────────────────────────

def _docx_heading(doc: Document, text: str, level: int, cfg: TemplateConfig) -> None:
    level = min(max(level, 1), 3)
    p = doc.add_heading(text, level=level)
    color = cfg.colors.secondary if level == 2 else cfg.colors.primary
    sizes = {1: 18, 2: 14, 3: 12}
    for run in p.runs:
        run.font.color.rgb = _rgb_color(color)
        run.font.size = Pt(sizes[level])
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    # Trait sous H1 pour le template rapport (style académique)
    if cfg.name == "rapport" and level == 1:
        _para_bottom_border(p, cfg.colors.secondary, sz=6)


def _docx_table(doc: Document, headers: list, rows: list, cfg: TemplateConfig) -> None:
    if not headers:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = str(h)
        _cell_bg(cell, cfg.colors.primary)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = _rgb_color(cfg.colors.text_light)
            run.font.bold = True
            run.font.size = Pt(10)

    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            if c_i < len(headers):
                cell = tbl.rows[r_i + 1].cells[c_i]
                cell.text = str(val)
                if r_i % 2 == 1:
                    _cell_bg(cell, cfg.colors.light)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)


def _docx_colored_section(doc: Document, block: dict, cfg: TemplateConfig) -> None:
    title = block.get("title", "")
    text  = block.get("text", "")
    if cfg.name in ("minimal", "contrat"):
        bg, fg = cfg.colors.light, cfg.colors.text_dark
    else:
        bg, fg = cfg.colors.secondary, cfg.colors.text_light

    if title:
        p = doc.add_paragraph()
        _para_bg(p, bg)
        run = p.add_run(title)
        run.font.bold  = True
        run.font.size  = Pt(13)
        run.font.color.rgb = _rgb_color(fg)
        p.paragraph_format.left_indent  = Pt(10)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)

    if text:
        p = doc.add_paragraph()
        _para_bg(p, bg)
        run = p.add_run(text)
        run.font.size  = Pt(11)
        run.font.color.rgb = _rgb_color(fg)
        p.paragraph_format.left_indent  = Pt(10)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(8)


def _docx_callout(doc: Document, block: dict, cfg: TemplateConfig) -> None:
    callout_type = block.get("callout_type", "note")
    ct    = CALLOUT_TYPES.get(callout_type, CALLOUT_TYPES["note"])
    label = block.get("label", "") or ct["label_default"]
    text  = block.get("text", "")

    p = doc.add_paragraph()
    _para_bg(p, ct["bg"])
    _para_left_border(p, ct["border"], sz=20)

    r = p.add_run(f"{label}  ")
    r.font.bold  = True
    r.font.size  = Pt(11)
    r.font.color.rgb = _rgb_color(ct["border"])

    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = _rgb_color(cfg.colors.text_dark)

    p.paragraph_format.left_indent  = Pt(14)
    p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)


def _docx_header_block(doc: Document, block: dict, cfg: TemplateConfig) -> None:
    title    = block.get("title", "")
    subtitle = block.get("subtitle", "")
    ref      = block.get("reference", "")
    date     = block.get("date", "")

    if title:
        if cfg.name == "commercial":
            p = doc.add_paragraph()
            _para_bg(p, cfg.colors.primary)
            run = p.add_run(title)
            run.font.bold  = True
            run.font.size  = Pt(22)
            run.font.color.rgb = _rgb_color(cfg.colors.text_light)
            p.paragraph_format.left_indent  = Pt(10)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
        else:
            p = doc.add_heading(title, level=1)
            for run in p.runs:
                run.font.color.rgb = _rgb_color(cfg.colors.primary)
                run.font.size = Pt(22)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            if cfg.name == "rapport":
                _para_bottom_border(p, cfg.colors.secondary)

    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size   = Pt(13)
        run.font.italic = True
        run.font.color.rgb = _rgb_color(cfg.colors.secondary)
        p.paragraph_format.space_after = Pt(4)

    if ref or date:
        p = doc.add_paragraph()
        run = p.add_run(" | ".join(filter(None, [ref, date])))
        run.font.size  = Pt(10)
        run.font.color.rgb = _rgb_color(cfg.colors.accent)
        p.paragraph_format.space_after = Pt(14)


# ─── DOCX renderer ────────────────────────────────────────────────────────────

def _render_docx(blocks: list[dict], cfg: TemplateConfig, company_name: str) -> bytes:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    mt, mb, ml, mr = cfg.page_margin_cm
    for sec in doc.sections:
        sec.top_margin    = Cm(mt)
        sec.bottom_margin = Cm(mb)
        sec.left_margin   = Cm(ml)
        sec.right_margin  = Cm(mr)

    _docx_page_header(doc, cfg, company_name)

    for block in blocks:
        t = block.get("type", "")

        if t == "cover_page":
            _docx_cover_page(doc, block, cfg, company_name)

        elif t == "heading":
            _docx_heading(doc, block.get("text", ""), block.get("level", 1), cfg)

        elif t == "paragraph":
            p = doc.add_paragraph(block.get("text", ""))
            p.paragraph_format.space_after = Pt(4)

        elif t == "bullet_list":
            sym = BULLET_SYMBOLS.get(block.get("style", cfg.bullet_default), "•")
            for item in block.get("items", []):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent       = Pt(24)
                p.paragraph_format.first_line_indent = Pt(-16)
                p.paragraph_format.space_after       = Pt(2)
                run = p.add_run(f"{sym}  {item}")
                run.font.size = Pt(11)

        elif t == "numbered_list":
            for idx, item in enumerate(block.get("items", []), start=1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent       = Pt(24)
                p.paragraph_format.first_line_indent = Pt(-16)
                p.paragraph_format.space_after       = Pt(2)
                run = p.add_run(f"{idx}.  {item}")
                run.font.size = Pt(11)

        elif t == "table":
            _docx_table(doc, block.get("headers", []), block.get("rows", []), cfg)
            doc.add_paragraph()

        elif t == "colored_section":
            _docx_colored_section(doc, block, cfg)

        elif t == "callout":
            _docx_callout(doc, block, cfg)

        elif t == "header_block":
            _docx_header_block(doc, block, cfg)

        elif t == "page_break":
            doc.add_page_break()

        elif t == "divider":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("─" * 50)
            r.font.size = Pt(10)
            r.font.color.rgb = _rgb_color(cfg.colors.accent)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── PDF helpers ───────────────────────────────────────────────────────────────

def _pdf_styles(cfg: TemplateConfig) -> dict:
    return {
        "h1": ParagraphStyle("h1",
            fontSize=18, fontName="Helvetica-Bold",
            textColor=_hc(cfg.colors.primary),
            spaceBefore=14, spaceAfter=4),
        "h2": ParagraphStyle("h2",
            fontSize=14, fontName="Helvetica-Bold",
            textColor=_hc(cfg.colors.secondary),
            spaceBefore=10, spaceAfter=4),
        "h3": ParagraphStyle("h3",
            fontSize=12, fontName="Helvetica-Bold",
            textColor=_hc(cfg.colors.primary),
            spaceBefore=8, spaceAfter=3),
        "body": ParagraphStyle("body",
            fontSize=10, fontName="Helvetica",
            textColor=_hc(cfg.colors.text_dark),
            spaceAfter=4, leading=14),
        "bullet": ParagraphStyle("bullet",
            fontSize=10, fontName="Helvetica",
            textColor=_hc(cfg.colors.text_dark),
            leftIndent=18, spaceAfter=2, leading=13),
        "header_title": ParagraphStyle("header_title",
            fontSize=22, fontName="Helvetica-Bold",
            textColor=_hc(cfg.colors.primary),
            spaceBefore=8, spaceAfter=4),
        "header_sub": ParagraphStyle("header_sub",
            fontSize=13, fontName="Helvetica-Oblique",
            textColor=_hc(cfg.colors.secondary),
            spaceAfter=4),
        "header_meta": ParagraphStyle("header_meta",
            fontSize=9, fontName="Helvetica",
            textColor=_hc(cfg.colors.accent),
            spaceAfter=14),
    }


def _pdf_page_callback(
    cfg: TemplateConfig, company_name: str, has_cover: bool = False
):
    def on_page(canvas, doc):
        canvas.saveState()
        w, h = A4

        # Pas d'en-tête sur la page de garde
        if has_cover and doc.page == 1:
            canvas.restoreState()
            return

        if cfg.has_page_header or company_name:
            if cfg.name in ("commercial", "rh"):
                r, g, b = _hex_rgb(cfg.colors.primary)
                canvas.setFillColorRGB(r / 255, g / 255, b / 255)
                canvas.rect(0, h - 28, w, 28, fill=1, stroke=0)
                canvas.setFillColorRGB(1, 1, 1)
                canvas.setFont("Helvetica-Bold", 10)
                label = company_name or (
                    "Boulga AI" if cfg.name == "commercial" else "Ressources Humaines"
                )
                canvas.drawString(20, h - 18, label)

            elif cfg.name == "rapport":
                r, g, b = _hex_rgb(cfg.colors.light)
                canvas.setFillColorRGB(r / 255, g / 255, b / 255)
                canvas.rect(0, h - 24, w, 24, fill=1, stroke=0)
                rs, gs, bs = _hex_rgb(cfg.colors.secondary)
                canvas.setFillColorRGB(rs / 255, gs / 255, bs / 255)
                canvas.rect(0, h - 26, w, 2, fill=1, stroke=0)
                rp, gp, bp = _hex_rgb(cfg.colors.primary)
                canvas.setFillColorRGB(rp / 255, gp / 255, bp / 255)
                canvas.setFont("Helvetica-Oblique", 9)
                canvas.drawString(20, h - 17, company_name or "Rapport")

            elif cfg.name == "contrat":
                ra, ga, ba = _hex_rgb(cfg.colors.accent)
                canvas.setFillColorRGB(ra / 255, ga / 255, ba / 255)
                canvas.setFont("Helvetica-Oblique", 9)
                label = company_name or "Document confidentiel"
                canvas.drawRightString(w - 20, h - 16, label)

            elif company_name:
                ra, ga, ba = _hex_rgb(cfg.colors.accent)
                canvas.setFillColorRGB(ra / 255, ga / 255, ba / 255)
                canvas.setFont("Helvetica", 9)
                canvas.drawRightString(w - 20, h - 16, company_name)

        # Numéro de page centré en bas
        ra, ga, ba = _hex_rgb(cfg.colors.accent)
        canvas.setFillColorRGB(ra / 255, ga / 255, ba / 255)
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(w / 2, 14, f"— {doc.page} —")
        canvas.restoreState()

    return on_page


# ─── PDF cover page ────────────────────────────────────────────────────────────

def _pdf_cover_page(
    block: dict, cfg: TemplateConfig, avail_w: float
) -> list:
    """Rendu PDF de la page de garde — retourne une liste de flowables."""
    title       = block.get("title", "Document")
    subtitle    = block.get("subtitle", "")
    institution = block.get("institution", "")
    author      = block.get("author", "")
    date        = block.get("date", "")
    ref         = block.get("reference", "")
    doc_type    = block.get("doc_type", "")
    meta        = [x for x in [author, date, ref] if x]

    def cpara(
        text: str, size: int, color: str,
        font: str = "Helvetica-Bold",
        align: int = TA_CENTER,
        sb: int = 0, sa: int = 8,
    ) -> Paragraph:
        return Paragraph(text, ParagraphStyle("cp",
            fontName=font, fontSize=size,
            textColor=_hc(color),
            alignment=align,
            spaceBefore=sb, spaceAfter=sa,
            leading=int(size * 1.3)))

    def band(text: str, bg: str, fg: str, size: int = 11,
             font: str = "Helvetica-Bold") -> Table:
        p = Paragraph(text, ParagraphStyle("band",
            fontName=font, fontSize=size, textColor=_hc(fg)))
        tbl = Table([[p]], colWidths=[avail_w])
        tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, -1), _hc(bg)),
            ("LEFTPADDING",   (0, 0), (-1, -1), 14),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 14),
            ("TOPPADDING",    (0, 0), (-1, -1), 12),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
        ]))
        return tbl

    story: list = []

    if cfg.name == "commercial":
        head = "  —  ".join(filter(None, [institution or "Document", doc_type]))
        story.append(band(head or "Document Commercial",
                          cfg.colors.primary, cfg.colors.text_light))
        story.append(Spacer(1, 160))
        story.append(cpara(title, 26, cfg.colors.primary, sa=10))
        if subtitle:
            story.append(cpara(subtitle, 14, cfg.colors.secondary,
                               font="Helvetica-Oblique", sa=6))
        story.append(Spacer(1, 160))
        story.append(HRFlowable(width="60%", thickness=2,
                                color=_hc(cfg.colors.secondary),
                                spaceBefore=0, spaceAfter=8))
        if meta:
            story.append(cpara("  ·  ".join(meta), 10, cfg.colors.accent,
                               font="Helvetica", sa=4))

    elif cfg.name == "rapport":
        if institution:
            story.append(band(institution, cfg.colors.light, cfg.colors.primary,
                              font="Helvetica"))
        story.append(HRFlowable(width="100%", thickness=2.5,
                                color=_hc(cfg.colors.secondary),
                                spaceBefore=0, spaceAfter=0))
        story.append(Spacer(1, 150))
        story.append(cpara(title, 24, cfg.colors.primary, sa=10))
        if subtitle:
            story.append(cpara(subtitle, 13, cfg.colors.secondary,
                               font="Helvetica-Oblique", sa=6))
        story.append(Spacer(1, 140))
        story.append(HRFlowable(width="100%", thickness=2.5,
                                color=_hc(cfg.colors.secondary),
                                spaceBefore=0, spaceAfter=8))
        if meta:
            story.append(cpara("  ·  ".join(meta), 10, cfg.colors.accent,
                               font="Helvetica", sa=4))

    elif cfg.name == "contrat":
        story.append(Spacer(1, 60))
        if doc_type:
            story.append(cpara(doc_type.upper(), 11, cfg.colors.accent,
                               sa=10))
        story.append(Spacer(1, 30))
        story.append(cpara(title, 22, cfg.colors.primary, sa=10))
        if subtitle:
            story.append(cpara(subtitle, 12, cfg.colors.accent,
                               font="Helvetica-Oblique", sa=6))
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width="80%", thickness=1.5,
                                color=_hc(cfg.colors.secondary),
                                spaceBefore=0, spaceAfter=16))
        if institution:
            story.append(cpara(institution, 11, cfg.colors.primary,
                               sa=4))
        story.append(Spacer(1, 100))
        if meta:
            story.append(cpara("  |  ".join(meta), 10, cfg.colors.accent,
                               font="Helvetica", sa=4))

    elif cfg.name == "rh":
        head = "  —  ".join(filter(None, [institution or "Ressources Humaines", doc_type]))
        story.append(band(head, cfg.colors.primary, cfg.colors.text_light, size=10))
        story.append(Spacer(1, 140))
        story.append(cpara(title, 22, cfg.colors.primary, sa=10))
        if subtitle:
            story.append(cpara(subtitle, 13, cfg.colors.secondary,
                               font="Helvetica-Oblique", sa=6))
        story.append(Spacer(1, 120))
        if author:
            story.append(cpara(author, 11, cfg.colors.primary, sa=4))
        meta2 = [x for x in [date, ref] if x]
        if meta2:
            story.append(cpara("  ·  ".join(meta2), 10, cfg.colors.accent,
                               font="Helvetica", sa=4))

    else:  # minimal
        story.append(Spacer(1, 200))
        story.append(cpara(title, 28, cfg.colors.primary, sa=12))
        if subtitle:
            story.append(cpara(subtitle, 14, cfg.colors.secondary,
                               font="Helvetica-Oblique", sa=6))
        story.append(Spacer(1, 16))
        story.append(HRFlowable(width="50%", thickness=1,
                                color=_hc(cfg.colors.accent),
                                spaceBefore=0, spaceAfter=10))
        if meta:
            story.append(cpara("  ·  ".join(meta), 10, cfg.colors.accent,
                               font="Helvetica", sa=4))

    story.append(PageBreak())
    return story


# ─── PDF block helpers ─────────────────────────────────────────────────────────

def _pdf_colored_section(
    block: dict, cfg: TemplateConfig, styles: dict, avail_w: float
) -> list:
    title = block.get("title", "")
    text  = block.get("text", "")

    if cfg.name in ("minimal", "contrat"):
        bg    = _hc(cfg.colors.light)
        t_col = cfg.colors.primary
        b_col = cfg.colors.text_dark
    else:
        bg    = _hc(cfg.colors.secondary)
        t_col = cfg.colors.text_light
        b_col = cfg.colors.text_light

    rows: list = []
    if title:
        rows.append(Paragraph(title, ParagraphStyle("cst",
            fontName="Helvetica-Bold", fontSize=12,
            textColor=_hc(t_col), leftIndent=10,
            spaceBefore=2, spaceAfter=2)))
    if text:
        rows.append(Paragraph(text, ParagraphStyle("csb",
            fontName="Helvetica", fontSize=10,
            textColor=_hc(b_col), leftIndent=10, leading=13)))
    if not rows:
        return []

    tbl = Table([[rows]], colWidths=[avail_w])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, -1), bg),
        ("LEFTPADDING",   (0, 0), (-1, -1), 12),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 12),
        ("TOPPADDING",    (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
    ]))
    return [tbl, Spacer(1, 8)]


def _pdf_callout(block: dict, cfg: TemplateConfig, avail_w: float) -> list:
    callout_type = block.get("callout_type", "note")
    ct    = CALLOUT_TYPES.get(callout_type, CALLOUT_TYPES["note"])
    label = block.get("label", "") or ct["label_default"]
    text  = block.get("text", "")

    full_text = (
        f'<font color="{ct["border"]}"><b>{label}</b></font>'
        f'<br/>{text}'
    )
    p = Paragraph(full_text, ParagraphStyle("callout_p",
        fontName="Helvetica", fontSize=10,
        textColor=_hc(cfg.colors.text_dark),
        leftIndent=8, rightIndent=4,
        leading=14, spaceBefore=0, spaceAfter=0))

    tbl = Table([[p]], colWidths=[avail_w])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, -1), _hc(ct["bg"])),
        ("LINEBEFORE",    (0, 0), (0, -1),  4, _hc(ct["border"])),
        ("LEFTPADDING",   (0, 0), (-1, -1), 14),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
        ("TOPPADDING",    (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    return [KeepTogether([tbl, Spacer(1, 8)])]


# ─── PDF renderer ──────────────────────────────────────────────────────────────

def _render_pdf(blocks: list[dict], cfg: TemplateConfig, company_name: str) -> bytes:
    buf = io.BytesIO()
    mt, mb, ml, mr = cfg.page_margin_cm
    has_cover  = bool(blocks) and blocks[0].get("type") == "cover_page"
    extra_top  = 1.0 if (cfg.has_page_header or company_name) else 0.0
    avail_w    = (21 - ml - mr) * cm

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=ml * cm, rightMargin=mr * cm,
        topMargin=(mt + extra_top) * cm, bottomMargin=(mb + 0.5) * cm,
    )

    styles  = _pdf_styles(cfg)
    on_page = _pdf_page_callback(cfg, company_name, has_cover)
    story: list = []

    for block in blocks:
        t = block.get("type", "")

        if t == "cover_page":
            story.extend(_pdf_cover_page(block, cfg, avail_w))

        elif t == "heading":
            level = min(max(block.get("level", 1), 1), 3)
            story.append(Paragraph(block.get("text", ""), styles[f"h{level}"]))
            # Trait sous H1 pour le template rapport (style académique)
            if cfg.name == "rapport" and level == 1:
                story.append(HRFlowable(
                    width="100%", thickness=1.5,
                    color=_hc(cfg.colors.secondary),
                    spaceBefore=2, spaceAfter=6,
                ))

        elif t == "paragraph":
            story.append(Paragraph(block.get("text", ""), styles["body"]))
            story.append(Spacer(1, 4))

        elif t == "bullet_list":
            sym = BULLET_SYMBOLS.get(block.get("style", cfg.bullet_default), "•")
            for item in block.get("items", []):
                story.append(Paragraph(f"{sym}  {item}", styles["bullet"]))
            story.append(Spacer(1, 4))

        elif t == "numbered_list":
            for i, item in enumerate(block.get("items", []), start=1):
                story.append(Paragraph(f"{i}.  {item}", styles["bullet"]))
            story.append(Spacer(1, 4))

        elif t == "table":
            headers = block.get("headers", [])
            rows    = block.get("rows", [])
            if not headers:
                continue

            th_sty = ParagraphStyle("th",
                fontName="Helvetica-Bold", fontSize=9,
                textColor=_hc(cfg.colors.text_light))
            td_sty = ParagraphStyle("td",
                fontName="Helvetica", fontSize=9,
                textColor=_hc(cfg.colors.text_dark), leading=12)

            col_w = avail_w / len(headers)
            data  = [[Paragraph(str(h), th_sty) for h in headers]]
            for row in rows:
                data.append([Paragraph(str(c), td_sty) for c in row])

            ts = [
                ("BACKGROUND",    (0, 0), (-1, 0),  _hc(cfg.colors.primary)),
                ("GRID",          (0, 0), (-1, -1), 0.5, grey),
                ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING",   (0, 0), (-1, -1), 6),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
                ("TOPPADDING",    (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
            for r_i in range(1, len(data)):
                if r_i % 2 == 0:
                    ts.append(("BACKGROUND", (0, r_i), (-1, r_i), _hc(cfg.colors.light)))

            tbl = Table(data, colWidths=[col_w] * len(headers))
            tbl.setStyle(TableStyle(ts))
            story.append(tbl)
            story.append(Spacer(1, 8))

        elif t == "colored_section":
            story.extend(_pdf_colored_section(block, cfg, styles, avail_w))

        elif t == "callout":
            story.extend(_pdf_callout(block, cfg, avail_w))

        elif t == "header_block":
            title    = block.get("title", "")
            subtitle = block.get("subtitle", "")
            ref      = block.get("reference", "")
            date     = block.get("date", "")

            if title:
                if cfg.name == "commercial":
                    tp = Paragraph(title, ParagraphStyle("ht",
                        fontSize=22, fontName="Helvetica-Bold",
                        textColor=_hc(cfg.colors.text_light),
                        leftIndent=10, spaceBefore=8, spaceAfter=4))
                    tbl = Table([[tp]], colWidths=[avail_w])
                    tbl.setStyle(TableStyle([
                        ("BACKGROUND",    (0, 0), (-1, -1), _hc(cfg.colors.primary)),
                        ("LEFTPADDING",   (0, 0), (-1, -1), 14),
                        ("TOPPADDING",    (0, 0), (-1, -1), 14),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 14),
                    ]))
                    story.append(tbl)
                    story.append(Spacer(1, 6))
                else:
                    story.append(Paragraph(title, styles["header_title"]))
                    if cfg.name == "rapport":
                        story.append(HRFlowable(
                            width="100%", thickness=1.5,
                            color=_hc(cfg.colors.secondary),
                            spaceBefore=2, spaceAfter=6,
                        ))

            if subtitle:
                story.append(Paragraph(subtitle, styles["header_sub"]))
            if ref or date:
                story.append(Paragraph(
                    " | ".join(filter(None, [ref, date])), styles["header_meta"]
                ))

        elif t == "page_break":
            story.append(PageBreak())

        elif t == "divider":
            story.append(HRFlowable(
                width="80%", thickness=1,
                color=_hc(cfg.colors.accent),
                spaceBefore=8, spaceAfter=8,
            ))

    doc.build(
        story or [Spacer(1, 1)],
        onFirstPage=on_page,
        onLaterPages=on_page,
    )
    return buf.getvalue()


# ─── Excel renderer ────────────────────────────────────────────────────────────

def _render_xlsx(blocks: list[dict], cfg: TemplateConfig, company_name: str) -> bytes:
    # Pre-pass : nombre max de colonnes (déterminé par le tableau le plus large)
    max_cols = max(
        (len(b.get("headers", [])) for b in blocks if b.get("type") == "table"),
        default=0,
    )
    max_cols = max(max_cols, 6)

    # Nom de l'onglet = titre du document (31 chars max pour Excel)
    sheet_title = ""
    for b in blocks[:2]:
        if b.get("type") in ("cover_page", "header_block"):
            sheet_title = b.get("title", "")[:31]
            break

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_title or "Document"

    # Largeur de colonne par défaut
    for c in range(1, max_cols + 1):
        ws.column_dimensions[get_column_letter(c)].width = 20

    # ── Helpers locaux ──────────────────────────────────────────────────────

    def _fill(hex_color: str) -> PatternFill:
        return PatternFill("solid", fgColor=hex_color.lstrip("#"))

    def _thin_border() -> Border:
        s = Side(style="thin", color="D0D0D0")
        return Border(left=s, right=s, top=s, bottom=s)

    def _write_row(
        row: int, value: str,
        bg: str = "", fg: str = "#1A1A1A",
        bold: bool = False, italic: bool = False,
        size: int = 10, height: int = 16,
        indent: int = 1, wrap: bool = False,
    ) -> None:
        """Écrit une ligne fusionnée sur max_cols colonnes."""
        cell = ws.cell(row=row, column=1, value=value)
        if bg:
            cell.fill = _fill(bg)
        cell.font = Font(size=size, bold=bold, italic=italic, color=fg.lstrip("#"))
        cell.alignment = Alignment(
            horizontal="left", vertical="center",
            indent=indent, wrap_text=wrap,
        )
        ws.row_dimensions[row].height = height
        if max_cols > 1:
            ws.merge_cells(
                start_row=row, start_column=1,
                end_row=row,   end_column=max_cols,
            )

    # ── Rendu bloc par bloc ─────────────────────────────────────────────────

    row = 1

    for block in blocks:
        t = block.get("type", "")

        if t == "cover_page":
            title       = block.get("title", "")
            subtitle    = block.get("subtitle", "")
            institution = block.get("institution", "") or company_name
            date        = block.get("date", "")
            ref         = block.get("reference", "")
            doc_type    = block.get("doc_type", "")

            if title:
                _write_row(row, title,
                           bg=cfg.colors.primary, fg=cfg.colors.text_light,
                           bold=True, size=18, height=40, indent=2)
                row += 1
            if subtitle:
                _write_row(row, subtitle,
                           bg=cfg.colors.secondary, fg=cfg.colors.text_light,
                           italic=True, size=12, height=26, indent=2)
                row += 1
            meta = [x for x in [institution, doc_type, date, ref] if x]
            if meta:
                _write_row(row, "  |  ".join(meta),
                           bg=cfg.colors.accent, fg=cfg.colors.text_light,
                           size=9, height=18, indent=2)
                row += 1
            row += 1  # ligne vide après la garde

        elif t == "header_block":
            title    = block.get("title", "")
            subtitle = block.get("subtitle", "")
            ref      = block.get("reference", "")
            date     = block.get("date", "")

            if title:
                _write_row(row, title,
                           bg=cfg.colors.primary, fg=cfg.colors.text_light,
                           bold=True, size=14, height=30, indent=2)
                row += 1
            if subtitle:
                _write_row(row, subtitle,
                           bg=cfg.colors.secondary, fg=cfg.colors.text_light,
                           italic=True, size=11, height=22, indent=2)
                row += 1
            if ref or date:
                _write_row(row, " | ".join(filter(None, [ref, date])),
                           fg=cfg.colors.accent, italic=True, size=9, height=16)
                row += 1
            row += 1

        elif t == "heading":
            level   = min(max(block.get("level", 1), 1), 3)
            text    = block.get("text", "")
            sizes   = {1: 13, 2: 11, 3: 10}
            heights = {1: 24, 2: 20, 3: 18}
            indents = {1: 1,  2: 2,  3: 3}
            bgs     = {1: cfg.colors.light, 2: "", 3: ""}
            fgs     = {1: cfg.colors.primary, 2: cfg.colors.secondary, 3: cfg.colors.primary}

            _write_row(row, text,
                       bg=bgs[level], fg=fgs[level], bold=True,
                       size=sizes[level], height=heights[level], indent=indents[level])

            if level == 1:
                rs, gs, bs = _hex_rgb(cfg.colors.secondary)
                border_color = f"{rs:02X}{gs:02X}{bs:02X}"
                for c in range(1, max_cols + 1):
                    existing = ws.cell(row=row, column=c).border
                    ws.cell(row=row, column=c).border = Border(
                        left=existing.left, right=existing.right, top=existing.top,
                        bottom=Side(style="medium", color=border_color),
                    )
            row += 1

        elif t == "paragraph":
            text  = block.get("text", "")
            lines = max(1, (len(text) + 99) // 100)
            _write_row(row, text, size=10, height=max(16, lines * 14 + 4),
                       indent=1, wrap=True)
            row += 1

        elif t == "bullet_list":
            sym = BULLET_SYMBOLS.get(block.get("style", cfg.bullet_default), "•")
            for item in block.get("items", []):
                _write_row(row, f"  {sym}  {item}", size=10, height=16, indent=2)
                row += 1

        elif t == "numbered_list":
            for idx, item in enumerate(block.get("items", []), start=1):
                _write_row(row, f"  {idx}.  {item}", size=10, height=16, indent=2)
                row += 1

        elif t == "table":
            headers   = block.get("headers", [])
            rows_data = block.get("rows", [])
            if not headers:
                continue

            n_cols = len(headers)
            thin   = _thin_border()

            rp, gp, bp = _hex_rgb(cfg.colors.primary)
            prim_fill = PatternFill("solid", fgColor=f"{rp:02X}{gp:02X}{bp:02X}")
            rl, gl, bl = _hex_rgb(cfg.colors.light)
            lite_fill  = PatternFill("solid", fgColor=f"{rl:02X}{gl:02X}{bl:02X}")

            # En-tête
            for c_i, header in enumerate(headers):
                cell = ws.cell(row=row, column=c_i + 1, value=header)
                cell.fill      = prim_fill
                cell.font      = Font(bold=True, size=10, color="FFFFFF")
                cell.alignment = Alignment(horizontal="center", vertical="center",
                                           wrap_text=True)
                cell.border    = thin
            ws.row_dimensions[row].height = 20
            row += 1

            # Lignes de données
            for r_i, data_row in enumerate(rows_data):
                for c_i, val in enumerate(data_row):
                    if c_i < n_cols:
                        cell = ws.cell(row=row, column=c_i + 1, value=val)
                        if r_i % 2 == 1:
                            cell.fill = lite_fill
                        cell.font      = Font(size=10)
                        cell.alignment = Alignment(horizontal="left", vertical="center",
                                                   wrap_text=True)
                        cell.border = thin
                ws.row_dimensions[row].height = 16
                row += 1
            row += 1  # ligne vide après le tableau

        elif t == "callout":
            callout_type = block.get("callout_type", "note")
            ct    = CALLOUT_TYPES.get(callout_type, CALLOUT_TYPES["note"])
            label = block.get("label", "") or ct["label_default"]
            text  = block.get("text", "")
            value = f"{label}: {text}" if label else text

            rb, gb, bb = _hex_rgb(ct["border"])
            border_hex = f"{rb:02X}{gb:02X}{bb:02X}"
            lines = max(1, (len(text) + 99) // 100)

            cell = ws.cell(row=row, column=1, value=value)
            cell.fill      = _fill(ct["bg"])
            cell.font      = Font(size=10, color=border_hex)
            cell.alignment = Alignment(horizontal="left", vertical="center",
                                       wrap_text=True, indent=2)
            cell.border    = Border(left=Side(style="thick", color=border_hex))
            ws.row_dimensions[row].height = max(18, lines * 14 + 6)
            if max_cols > 1:
                ws.merge_cells(start_row=row, start_column=1,
                               end_row=row, end_column=max_cols)
            row += 1

        elif t == "colored_section":
            title = block.get("title", "")
            text  = block.get("text", "")
            value = f"{title}  —  {text}" if (title and text) else (title or text)
            lines = max(1, (len(value) + 99) // 100)
            _write_row(row, value,
                       bg=cfg.colors.secondary, fg=cfg.colors.text_light,
                       bold=bool(title), size=10,
                       height=max(18, lines * 14 + 6), indent=2, wrap=True)
            row += 1

        elif t == "divider":
            ra, ga, ba = _hex_rgb(cfg.colors.accent)
            accent_hex = f"{ra:02X}{ga:02X}{ba:02X}"
            for c in range(1, max_cols + 1):
                ws.cell(row=row, column=c).border = Border(
                    bottom=Side(style="thin", color=accent_hex)
                )
            ws.row_dimensions[row].height = 6
            row += 1

        elif t == "page_break":
            ws.row_dimensions[row].height = 8
            row += 2

    # ── Auto-ajustement des largeurs de colonnes ────────────────────────────
    # Les cellules fusionnées (blocs texte) ont value=None sur col 2..N → ignorées
    for col_idx in range(1, max_cols + 1):
        col_letter = get_column_letter(col_idx)
        max_len = 20  # minimum
        for cell in ws[col_letter]:
            if cell.value is not None:
                max_len = max(max_len, min(len(str(cell.value)) + 3, 55))
        ws.column_dimensions[col_letter].width = max_len

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
